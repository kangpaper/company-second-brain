from io import BytesIO

import pymupdf
import pytest
from docx import Document
from openpyxl import Workbook

from company_brain.ingestion.parsers import ParseError, parse_content


def make_xlsx() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Customers"
    sheet.append(["name", "revenue"])
    sheet.append(["Acme", 1200])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Quarterly Review", level=1)
    document.add_paragraph("Revenue increased by 12 percent.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Customer"
    table.cell(0, 1).text = "Risk"
    table.cell(1, 0).text = "Acme"
    table.cell(1, 1).text = "Low"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_table_only_docx(headers: list[str], values: list[str]) -> bytes:
    document = Document()
    table = document.add_table(rows=2, cols=len(headers))
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        table.cell(1, index).text = values[index]
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf(text: str = "Invoice INV-100 total 500") -> bytes:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    result = document.tobytes()
    document.close()
    return result


@pytest.mark.parametrize(
    ("media_type", "filename", "content", "expected", "candidate_kind"),
    [
        (
            "text/markdown",
            "review.md",
            b"# Customer Review\n\nAcme is healthy.",
            "Acme is healthy",
            "section",
        ),
        (
            "text/plain",
            "notes.txt",
            b"Acme plain text note.",
            "Acme plain text note",
            "section",
        ),
        (
            "text/csv",
            "customers.csv",
            b"name,revenue\nAcme,1200\n",
            "Acme",
            "row",
        ),
        (
            "text/html",
            "review.html",
            (
                b"<html><head><script>ignore()</script></head><body>"
                b"<h1>Review</h1><p>Acme healthy</p></body></html>"
            ),
            "Acme healthy",
            "section",
        ),
        (
            "text/html",
            "https://example.com/review",
            b"<article><h1>Web Review</h1><p>Acme web source</p></article>",
            "Acme web source",
            "section",
        ),
    ],
)
def test_text_fixture_parsers_preserve_candidate_provenance(
    media_type: str,
    filename: str,
    content: bytes,
    expected: str,
    candidate_kind: str,
) -> None:
    parsed = parse_content(media_type, filename, content)

    assert expected in parsed.text
    assert parsed.candidates
    assert parsed.candidates[0].kind == candidate_kind
    assert parsed.candidates[0].locator
    assert parsed.metadata["byte_size"] == len(content)


def test_xlsx_fixture_extracts_rows_with_sheet_and_row_locator() -> None:
    parsed = parse_content(
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "customers.xlsx",
        make_xlsx(),
    )

    assert "Acme" in parsed.text
    assert parsed.candidates[0].data == {"name": "Acme", "revenue": 1200}
    assert parsed.candidates[0].locator == {"sheet": "Customers", "row": 2}


def test_docx_fixture_extracts_paragraphs_and_table_rows() -> None:
    parsed = parse_content(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "review.docx",
        make_docx(),
    )

    assert "Revenue increased" in parsed.text
    assert {candidate.kind for candidate in parsed.candidates} == {"paragraph", "table_row"}
    assert any(candidate.locator == {"table": 1, "row": 2} for candidate in parsed.candidates)


def test_docx_blank_table_fails_without_fabricated_text() -> None:
    with pytest.raises(ParseError, match="no extractable text"):
        parse_content(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "blank-table.docx",
            make_table_only_docx(["", ""], ["", ""]),
        )


def test_docx_preserves_columns_with_blank_or_duplicate_headers() -> None:
    parsed = parse_content(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "headers.docx",
        make_table_only_docx(["", "Risk", "Risk"], ["Acme", "Low", "Monitored"]),
    )

    assert parsed.candidates[0].data == {
        "column_1": "Acme",
        "Risk": "Low",
        "Risk_2": "Monitored",
    }


@pytest.mark.parametrize(
    ("headers", "values", "expected"),
    [
        (
            ["Risk", "Risk", "Risk_2"],
            ["A", "B", "C"],
            {"Risk": "A", "Risk_2": "B", "Risk_2_2": "C"},
        ),
        (
            ["", "column_1"],
            ["A", "B"],
            {"column_1": "A", "column_1_2": "B"},
        ),
        (
            ["Risk", "Risk", "Risk_2", "Risk", "Risk_3"],
            ["A", "B", "C", "D", "E"],
            {
                "Risk": "A",
                "Risk_2": "B",
                "Risk_2_2": "C",
                "Risk_3": "D",
                "Risk_3_2": "E",
            },
        ),
    ],
)
def test_docx_header_normalization_is_globally_collision_safe(
    headers: list[str], values: list[str], expected: dict[str, str]
) -> None:
    parsed = parse_content(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "collision.docx",
        make_table_only_docx(headers, values),
    )

    assert parsed.candidates[0].data == expected
    assert list(parsed.candidates[0].data.values()) == values


def test_pdf_fixture_extracts_page_text_with_page_locator() -> None:
    parsed = parse_content("application/pdf", "invoice.pdf", make_pdf())

    assert "INV-100" in parsed.text
    assert parsed.candidates[0].locator == {"page": 1}


def test_malformed_and_scanned_pdf_fail_without_fabricated_text() -> None:
    blank = pymupdf.open()
    blank.new_page()
    blank_bytes = blank.tobytes()
    blank.close()

    with pytest.raises(ParseError, match="Invalid PDF"):
        parse_content("application/pdf", "broken.pdf", b"not a pdf")
    with pytest.raises(ParseError, match="no extractable text"):
        parse_content("application/pdf", "scan.pdf", blank_bytes)


def test_parser_rejects_unsupported_media_type_and_oversized_input() -> None:
    with pytest.raises(ParseError, match="Unsupported media type"):
        parse_content("application/octet-stream", "data.bin", b"abc")
    with pytest.raises(ParseError, match="exceeds"):
        parse_content("text/markdown", "huge.md", b"x" * (10 * 1024 * 1024 + 1))


def test_parser_stops_when_candidate_limit_is_exceeded() -> None:
    rows = "name\n" + "\n".join(f"Customer {index}" for index in range(5001))

    with pytest.raises(ParseError, match="candidate count exceeds 5000"):
        parse_content("text/csv", "too-many.csv", rows.encode())
